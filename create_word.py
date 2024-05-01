import os
from faker import Faker
from docx import Document

# 创建一个Faker实例
fake = Faker()

# 创建一个Document实例
doc = Document()

# 添加标题
doc.add_heading('工作报告', 0)

# 添加一些员工和日期
employees = [fake.name() for _ in range(10)]
dates = [fake.date_between(start_date='-1y', end_date='today') for _ in range(10)]

# 添加工作报告内容
for i in range(10):
    doc.add_heading('员工: ' + employees[i], level=1)
    doc.add_paragraph('日期: ' + str(dates[i]))
    doc.add_paragraph('工作内容: ')
    doc.add_paragraph(fake.text())

# 创建文件路径
file_path = os.path.join('data', '工作报告.docx')

# 将Document保存为Word文件
doc.save(file_path)
